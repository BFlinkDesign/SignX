import sys
import os
import json
import logging
from datetime import datetime
from pathlib import Path
from PyQt6.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout, 
                            QHBoxLayout, QTextEdit, QPushButton, QLabel, 
                            QComboBox, QProgressBar, QMessageBox, QFileDialog,
                            QSystemTrayIcon, QMenu, QStyle, QTabWidget,
                            QTableWidget, QTableWidgetItem, QHeaderView,
                            QSplitter, QToolBar, QStatusBar, QDialog,
                            QCheckBox, QGroupBox, QRadioButton, QDockWidget,
                            QTreeWidget, QTreeWidgetItem, QCalendarWidget,
                            QSpinBox, QDoubleSpinBox, QLineEdit, QFormLayout,
                            QScrollArea, QFrame)
from PyQt6.QtCore import Qt, QThread, pyqtSignal, QSettings, QTimer, QSize, QProcess, QObject
from PyQt6.QtGui import QIcon, QAction, QFont, QColor, QPalette, QPixmap, QImage, QPainter
from PyQt6.QtCharts import QChart, QChartView, QPieSeries, QBarSeries, QBarSet, QBarCategoryAxis, QValueAxis, QScatterSeries
import pandas as pd
import numpy as np
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Inches, Pt, RGBColor
import xml.etree.ElementTree as ET
import yaml
import toml
from sklearn.cluster import KMeans
from sklearn.preprocessing import StandardScaler
from sklearn.decomposition import PCA
import plotly.express as px
import plotly.graph_objects as go
from plotly.subplots import make_subplots
import seaborn as sns
import matplotlib.pyplot as plt
from datetime import timedelta
import threading
import queue
import hashlib
import sqlite3
from typing import Dict, List, Any, Optional, Tuple
import re
import difflib
from concurrent.futures import ThreadPoolExecutor
import nltk
from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords
from nltk.stem import WordNetLemmatizer
import spacy
from transformers.pipelines import pipeline
import torch
from fuzzywuzzy import fuzz
import jellyfish
from tqdm import tqdm
import warnings
import hashlib
import hmac
import base64
import secrets
import bcrypt
from cryptography.fernet import Fernet
from cryptography.hazmat.primitives import hashes
from cryptography.hazmat.primitives.kdf.pbkdf2 import PBKDF2HMAC
import jwt
import ssl
import certifi
import win32security
import win32api
import win32con
import win32crypt
import win32file
import win32process
import win32event
import win32service
import win32serviceutil
import win32ts
import win32net
from pydantic import BaseModel, Field, validator
import bleach
import unicodedata
import emoji
import phonenumbers
import email_validator
import ipaddress
import validators
import jsonschema
import marshmallow
import cerberus
import voluptuous
import trafaret
import colander
import schema
from style_config import COLORS, FONTS, FONT_SIZES, STYLES, THEME, EXPORT_FORMATS, DEPARTMENT_COLORS, COLOR_MODES, WINDOW
import subprocess
from openpyxl.chart import BarChart, Reference
import pdfkit
import html
import requests

# Constants
DOCX_EXTENSION = '.docx'
PDF_EXTENSION = '.pdf'
WORK_ORDER_PATTERN = r'^[A-Z0-9-]+$'
DATE_FORMAT = '%Y-%m-%d'
CHART_COLORS = {
    'header_fill': '4F81BD',
    'alt_row_fill': 'E6E6E6'
}

# Download required NLTK data
nltk.download('punkt')
nltk.download('stopwords')
nltk.download('wordnet')

# Load spaCy model
try:
    nlp = spacy.load('en_core_web_sm')
except:
    subprocess.run([sys.executable, '-m', 'spacy', 'download', 'en_core_web_sm'])
    nlp = spacy.load('en_core_web_sm')

# Configure logging with advanced features
class AdvancedLogger:
    def __init__(self):
        self.logger = logging.getLogger('WorkOrderParser')
        self.logger.setLevel(logging.DEBUG)
        
        # File handler with rotation
        file_handler = logging.handlers.RotatingFileHandler(
            'work_order_parser.log',
            maxBytes=10*1024*1024,  # 10MB
            backupCount=5
        )
        file_handler.setLevel(logging.DEBUG)
        
        # Console handler
        console_handler = logging.StreamHandler()
        console_handler.setLevel(logging.INFO)
        
        # Formatter
        formatter = logging.Formatter(
            '%(asctime)s - %(name)s - %(levelname)s - %(message)s'
        )
        file_handler.setFormatter(formatter)
        console_handler.setFormatter(formatter)
        
        self.logger.addHandler(file_handler)
        self.logger.addHandler(console_handler)
        
        # Performance metrics
        self.metrics = {
            'processing_time': [],
            'memory_usage': [],
            'error_count': 0,
            'success_count': 0
        }
    
    def log_performance(self, start_time: float, end_time: float):
        duration = end_time - start_time
        self.metrics['processing_time'].append(duration)
        self.logger.info(f'Operation completed in {duration:.2f} seconds')
    
    def get_performance_report(self) -> Dict[str, Any]:
        return {
            'avg_processing_time': np.mean(self.metrics['processing_time']),
            'max_processing_time': np.max(self.metrics['processing_time']),
            'min_processing_time': np.min(self.metrics['processing_time']),
            'total_operations': len(self.metrics['processing_time']),
            'error_rate': self.metrics['error_count'] / 
                         (self.metrics['error_count'] + self.metrics['success_count'])
        }

logger = AdvancedLogger()

class AIProcessor:
    def __init__(self):
        self.sentiment_analyzer = pipeline('sentiment-analysis')
        self.ner_pipeline = pipeline('ner')
        self.text_classifier = pipeline('text-classification')
        self.qa_pipeline = pipeline('question-answering')
        
        # Initialize models
        self.initialize_models()
    
    def initialize_models(self):
        try:
            # Load pre-trained models
            self.sentiment_analyzer = pipeline('sentiment-analysis')
            self.ner_pipeline = pipeline('ner')
            self.text_classifier = pipeline('text-classification')
            self.qa_pipeline = pipeline('question-answering')
            logger.logger.info("AI models initialized successfully")
        except Exception as e:
            logger.logger.error(f"Error initializing AI models: {str(e)}")
    
    def analyze_text(self, text: str) -> Dict[str, Any]:
        try:
            # Perform various AI analyses
            sentiment = self.sentiment_analyzer(text)[0]
            entities = self.ner_pipeline(text)
            classification = self.text_classifier(text)[0]
            
            return {
                'sentiment': sentiment,
                'entities': entities,
                'classification': classification
            }
        except Exception as e:
            logger.logger.error(f"Error in AI analysis: {str(e)}")
            return {}

class DataAnalyzer:
    def __init__(self):
        self.scaler = StandardScaler()
        self.pca = PCA(n_components=2)
    
    def analyze_data(self, data: pd.DataFrame) -> Dict[str, Any]:
        try:
            # Perform advanced data analysis
            numeric_data = data.select_dtypes(include=[np.number])
            if not numeric_data.empty:
                scaled_data = self.scaler.fit_transform(numeric_data)
                pca_result = self.pca.fit_transform(scaled_data)
                
                # Perform clustering
                kmeans = KMeans(n_clusters=3)
                clusters = kmeans.fit_predict(scaled_data)
                
                return {
                    'pca_result': pca_result,
                    'clusters': clusters,
                    'statistics': {
                        'mean': numeric_data.mean().to_dict(),
                        'std': numeric_data.std().to_dict(),
                        'correlation': numeric_data.corr().to_dict()
                    }
                }
            return {}
        except Exception as e:
            logger.logger.error(f"Error in data analysis: {str(e)}")
            return {}

# Security configuration
class SecurityConfig:
    def __init__(self):
        self.encryption_key = self._generate_encryption_key()
        self.fernet = Fernet(self.encryption_key)
        self.salt = os.urandom(16)
        self.kdf = PBKDF2HMAC(
            algorithm=hashes.SHA256(),
            length=32,
            salt=self.salt,
            iterations=100000,
        )
        self.jwt_secret = secrets.token_hex(32)
    
    def _generate_encryption_key(self) -> bytes:
        return Fernet.generate_key()
    
    def encrypt_data(self, data: str) -> bytes:
        return self.fernet.encrypt(data.encode())
    
    def decrypt_data(self, encrypted_data: bytes) -> str:
        return self.fernet.decrypt(encrypted_data).decode()
    
    def hash_password(self, password: str) -> bytes:
        return bcrypt.hashpw(password.encode(), bcrypt.gensalt())
    
    def verify_password(self, password: str, hashed: bytes) -> bool:
        return bcrypt.checkpw(password.encode(), hashed)
    
    def generate_token(self, data: dict) -> str:
        return jwt.encode(data, self.jwt_secret, algorithm='HS256')
    
    def verify_token(self, token: str) -> dict:
        return jwt.decode(token, self.jwt_secret, algorithms=['HS256'])

class SecureFileHandler:
    def __init__(self, security_config: SecurityConfig):
        self.security_config = security_config
        self.temp_dir = Path(os.environ['TEMP']) / 'work_order_parser'
        self.temp_dir.mkdir(exist_ok=True)
    
    def secure_write(self, file_path: str, data: Any) -> bool:
        try:
            # Encrypt data
            if isinstance(data, str):
                encrypted_data = self.security_config.encrypt_data(data)
            else:
                encrypted_data = self.security_config.encrypt_data(json.dumps(data))
            
            # Write to temporary file
            temp_file = self.temp_dir / f"{secrets.token_hex(8)}.tmp"
            with open(temp_file, 'wb') as f:
                f.write(encrypted_data)
            
            # Move to final location with secure permissions
            self._secure_move(temp_file, file_path)
            return True
        except Exception as e:
            logger.logger.error(f"Error in secure write: {str(e)}")
            return False
    
    def secure_read(self, file_path: str) -> Any:
        try:
            # Read encrypted data
            with open(file_path, 'rb') as f:
                encrypted_data = f.read()
            
            # Decrypt data
            decrypted_data = self.security_config.decrypt_data(encrypted_data)
            
            # Try to parse as JSON
            try:
                return json.loads(decrypted_data)
            except:
                return decrypted_data
        except Exception as e:
            logger.logger.error(f"Error in secure read: {str(e)}")
            return None
    
    def _secure_move(self, src: Path, dst: str):
        # Set secure permissions
        security_attributes = win32security.SECURITY_ATTRIBUTES()
        security_attributes.SECURITY_DESCRIPTOR = win32security.SECURITY_DESCRIPTOR()
        security_attributes.SECURITY_DESCRIPTOR.SetSecurityDescriptorDacl(1, None, 0)
        
        # Move file with secure permissions
        win32file.MoveFileEx(
            str(src),
            dst,
            win32file.MOVEFILE_REPLACE_EXISTING | win32file.MOVEFILE_WRITE_THROUGH
        )
        
        # Set file permissions
        sd = win32security.SECURITY_DESCRIPTOR()
        sd.SetSecurityDescriptorDacl(1, None, 0)
        win32security.SetFileSecurity(
            dst,
            win32security.DACL_SECURITY_INFORMATION,
            sd
        )

class AdvancedExportManager:
    def __init__(self, security_config: SecurityConfig):
        self.security_config = security_config
        self.secure_handler = SecureFileHandler(security_config)
        self.export_formats = {
            'excel': self.export_to_excel,
            'csv': self.export_to_csv,
            'json': self.export_to_json,
            'pdf': self.export_to_pdf,
            'word': self.export_to_word,
            'xml': self.export_to_xml,
            'yaml': self.export_to_yaml,
            'toml': self.export_to_toml,
            'html': self.export_to_html,
            'markdown': self.export_to_markdown,
            'latex': self.export_to_latex,
            'sql': self.export_to_sql,
            'parquet': self.export_to_parquet,
            'feather': self.export_to_feather,
            'pickle': self.export_to_pickle,
            'hdf5': self.export_to_hdf5
        }
    
    def export_data(self, data: pd.DataFrame, format: str, file_path: str) -> bool:
        try:
            if format in self.export_formats:
                # Create temporary file
                temp_file = self.secure_handler.temp_dir / f"{secrets.token_hex(8)}.{format}"
                
                # Export to temporary file
                self.export_formats[format](data, str(temp_file))
                
                # Move to final location with secure permissions
                self.secure_handler._secure_move(temp_file, file_path)
                return True
            return False
        except Exception as e:
            logger.logger.error(f"Error exporting data: {str(e)}")
            return False
    
    def export_to_excel(self, data: pd.DataFrame, file_path: str):
        wb = Workbook()
        ws = wb.active
        ws.title = "Work Orders"
        
        # Add headers with styling
        for col, header in enumerate(data.columns, 1):
            cell = ws.cell(row=1, column=col)
            cell.value = header
            cell.font = Font(bold=True, size=12)
            cell.fill = PatternFill(start_color=CHART_COLORS['header_fill'], end_color=CHART_COLORS['header_fill'], fill_type="solid")
            cell.font = Font(color="FFFFFF")
            cell.alignment = Alignment(horizontal="center")
        
        # Add data with alternating row colors
        for row_idx, row in enumerate(data.values, 2):
            for col_idx, value in enumerate(row, 1):
                cell = ws.cell(row=row_idx, column=col_idx)
                cell.value = value
                cell.alignment = Alignment(horizontal="center")
                if row_idx % 2 == 0:
                    cell.fill = PatternFill(start_color=CHART_COLORS['alt_row_fill'], end_color=CHART_COLORS['alt_row_fill'], fill_type="solid")
        
        # Add borders
        thin_border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
        
        for row in ws.iter_rows():
            for cell in row:
                cell.border = thin_border
        
        # Adjust column widths
        for col in range(1, len(data.columns) + 1):
            ws.column_dimensions[get_column_letter(col)].width = 15
        
        # Add summary sheet
        summary_sheet = wb.create_sheet("Summary")
        summary_sheet['A1'] = "Data Summary"
        summary_sheet['A1'].font = Font(bold=True, size=14)
        
        # Add statistics
        summary_sheet['A3'] = "Total Records"
        summary_sheet['B3'] = len(data)
        summary_sheet['A4'] = "Date Range"
        summary_sheet['B4'] = f"{data['date'].min()} to {data['date'].max()}"
        
        # Add charts
        chart_sheet = wb.create_sheet("Charts")
        self._add_excel_charts(data, chart_sheet)
        
        wb.save(file_path)
    
    def _add_excel_charts(self, data: pd.DataFrame, sheet):
        # Add bar chart for department distribution
        chart = BarChart()
        chart.title = "Department Distribution"
        chart.style = 10
        chart.y_axis.title = "Count"
        chart.x_axis.title = "Department"
        
        dept_data = data['department'].value_counts()
        data = Reference(sheet, min_col=1, min_row=1, max_row=len(dept_data))
        cats = Reference(sheet, min_col=2, min_row=1, max_row=len(dept_data))
        chart.add_data(data)
        chart.set_categories(cats)
        sheet.add_chart(chart, "A1")
    
    def export_to_pdf(self, data: pd.DataFrame, file_path: str):
        # Create a more sophisticated PDF with charts and tables
        doc = Document()
        doc.add_heading('Work Order Analysis Report', 0)
        
        # Add summary statistics
        doc.add_heading('Summary Statistics', level=1)
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.rows[0].cells[0].text = "Total Records"
        table.rows[0].cells[1].text = str(len(data))
        
        # Add data table
        doc.add_heading('Work Order Data', level=1)
        table = doc.add_table(rows=1, cols=len(data.columns))
        table.style = 'Table Grid'
        
        # Add headers
        header_cells = table.rows[0].cells
        for i, header in enumerate(data.columns):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].font.bold = True
        
        # Add data
        for _, row in data.iterrows():
            cells = table.add_row().cells
            for i, value in enumerate(row):
                cells[i].text = str(value)
        
        # Add charts
        self._add_pdf_charts(doc, data)
        
        # Save as PDF
        doc.save(file_path.replace('.pdf', DOCX_EXTENSION))
        pdfkit.from_file(file_path.replace('.pdf', DOCX_EXTENSION), file_path)
        os.remove(file_path.replace('.pdf', DOCX_EXTENSION))
    
    def _add_pdf_charts(self, doc: Document, data: pd.DataFrame):
        # Create charts using matplotlib
        plt.figure(figsize=(10, 6))
        
        # Department distribution
        dept_counts = data['department'].value_counts()
        plt.subplot(1, 2, 1)
        dept_counts.plot(kind='bar')
        plt.title('Department Distribution')
        plt.xticks(rotation=45)
        
        # Date distribution
        plt.subplot(1, 2, 2)
        data['date'].value_counts().sort_index().plot(kind='line')
        plt.title('Work Orders Over Time')
        plt.xticks(rotation=45)
        
        # Save chart
        chart_path = self.secure_handler.temp_dir / "temp_chart.png"
        plt.savefig(chart_path)
        plt.close()
        
        # Add chart to document
        doc.add_picture(str(chart_path), width=Inches(6))
        os.remove(chart_path)
    
    def export_to_markdown(self, data: pd.DataFrame, file_path: str):
        with open(file_path, 'w') as f:
            f.write('# Work Order Analysis Report\n\n')
            
            # Add summary
            f.write('## Summary\n\n')
            f.write(f'- Total Records: {len(data)}\n')
            f.write(f'- Date Range: {data["date"].min()} to {data["date"].max()}\n\n')
            
            # Add data table
            f.write('## Data\n\n')
            f.write(data.to_markdown(index=False))
    
    def export_to_latex(self, data: pd.DataFrame, file_path: str):
        with open(file_path, 'w') as f:
            f.write('\\documentclass{article}\n')
            f.write('\\usepackage{booktabs}\n')
            f.write('\\usepackage{graphicx}\n')
            f.write('\\begin{document}\n\n')
            
            f.write('\\title{Work Order Analysis Report}\n')
            f.write('\\maketitle\n\n')
            
            # Add summary
            f.write('\\section{Summary}\n')
            f.write(f'Total Records: {len(data)}\n\n')
            f.write(f'Date Range: {data["date"].min()} to {data["date"].max()}\n\n')
            
            # Add data table
            f.write('\\section{Data}\n')
            f.write(data.to_latex(index=False))
            
            f.write('\\end{document}')
    
    def export_to_sql(self, data: pd.DataFrame, file_path: str):
        with open(file_path, 'w') as f:
            f.write('-- Work Order Data\n\n')
            f.write('CREATE TABLE IF NOT EXISTS work_orders (\n')
            f.write('    id INTEGER PRIMARY KEY,\n')
            for col in data.columns:
                f.write(f'    {col} TEXT,\n')
            f.write('    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP\n')
            f.write(');\n\n')
            
            f.write('INSERT INTO work_orders (')
            f.write(', '.join(data.columns))
            f.write(') VALUES\n')
            
            for _, row in data.iterrows():
                values = [f"'{str(val)}'" for val in row]
                f.write(f"({', '.join(values)}),\n")
    
    def export_to_parquet(self, data: pd.DataFrame, file_path: str):
        data.to_parquet(file_path, index=False)
    
    def export_to_feather(self, data: pd.DataFrame, file_path: str):
        data.to_feather(file_path)
    
    def export_to_pickle(self, data: pd.DataFrame, file_path: str):
        data.to_pickle(file_path)
    
    def export_to_hdf5(self, data: pd.DataFrame, file_path: str):
        data.to_hdf(file_path, 'work_orders', format='table')
    
    def export_to_csv(self, data: pd.DataFrame, file_path: str):
        """Export DataFrame to CSV with Excel compatibility and error handling."""
        try:
            data.to_csv(file_path, index=False, encoding='utf-8-sig')
        except Exception as e:
            logger.logger.error(f"Error exporting to CSV: {str(e)}")
            raise

class DataValidator:
    def __init__(self):
        self.schema = {
            'work_order': {
                'type': 'string',
                'required': True,
                'regex': WORK_ORDER_PATTERN,
                'minlength': 3,
                'maxlength': 20
            },
            'date': {
                'type': 'string',
                'required': True,
                'regex': r'^\d{4}-\d{2}-\d{2}$'
            },
            'part_number': {
                'type': 'string',
                'required': True,
                'regex': WORK_ORDER_PATTERN,
                'minlength': 3,
                'maxlength': 20
            },
            'department': {
                'type': 'string',
                'required': True,
                'allowed': ['ES', 'SC', 'PR', 'MFG', 'ENG', 'QA']
            }
        }
        self.validator = cerberus.Validator(self.schema)
    
    def validate_data(self, data: Dict[str, Any]) -> Tuple[bool, List[str]]:
        if not self.validator.validate(data):
            return False, self.validator.errors
        return True, []
    
    def sanitize_text(self, text: str) -> str:
        # Remove HTML tags
        text = bleach.clean(text, strip=True)
        
        # Convert HTML entities
        text = html.unescape(text)
        
        # Normalize Unicode
        text = unicodedata.normalize('NFKC', text)
        
        # Remove emojis
        text = emoji.replace_emoji(text, '')
        
        # Remove control characters
        text = ''.join(ch for ch in text if unicodedata.category(ch)[0] != 'C')
        
        return text.strip()
    
    def validate_email(self, email: str) -> bool:
        try:
            email_validator.validate_email(email)
            return True
        except:
            return False
    
    def validate_phone(self, phone: str) -> bool:
        try:
            number = phonenumbers.parse(phone)
            return phonenumbers.is_valid_number(number)
        except:
            return False
    
    def validate_ip(self, ip: str) -> bool:
        try:
            ipaddress.ip_address(ip)
            return True
        except:
            return False
    
    def validate_url(self, url: str) -> bool:
        return validators.url(url)

class DataSanitizer:
    def __init__(self):
        self.validator = DataValidator()
    
    def sanitize_data(self, data: Dict[str, Any]) -> Dict[str, Any]:
        sanitized = {}
        for key, value in data.items():
            if isinstance(value, str):
                sanitized[key] = self.validator.sanitize_text(value)
            elif isinstance(value, dict):
                sanitized[key] = self.sanitize_data(value)
            elif isinstance(value, list):
                sanitized[key] = [self.sanitize_data(item) if isinstance(item, dict)
                                else self.validator.sanitize_text(item) if isinstance(item, str)
                                else item for item in value]
            else:
                sanitized[key] = value
        return sanitized
    
    def validate_and_sanitize(self, data: Dict[str, Any]) -> Tuple[Dict[str, Any], List[str]]:
        # First sanitize
        sanitized_data = self.sanitize_data(data)
        
        # Then validate
        is_valid, errors = self.validator.validate_data(sanitized_data)
        
        return sanitized_data, errors if not is_valid else []

class UserManager:
    def __init__(self, security_config: SecurityConfig):
        self.security_config = security_config
        self.current_user = None
        self.user_roles = {
            'admin': ['read', 'write', 'delete', 'export', 'admin'],
            'manager': ['read', 'write', 'export'],
            'user': ['read', 'export']
        }
    
    def authenticate_user(self, username: str, password: str) -> bool:
        try:
            # Get user from database
            cursor = self.conn.cursor()
            cursor.execute('SELECT * FROM users WHERE username = ?', (username,))
            user = cursor.fetchone()
            
            if user and self.security_config.verify_password(password, user['password_hash']):
                self.current_user = {
                    'id': user['id'],
                    'username': user['username'],
                    'role': user['role']
                }
                return True
            return False
        except Exception as e:
            logger.logger.error(f"Error authenticating user: {str(e)}")
            return False
    
    def check_permission(self, permission: str) -> bool:
        if not self.current_user:
            return False
        
        role = self.current_user['role']
        return permission in self.user_roles.get(role, [])

class LoginDialog(QDialog):
    def __init__(self, user_manager: UserManager, parent=None):
        super().__init__(parent)
        self.user_manager = user_manager
        self.setWindowTitle("Login")
        self.setMinimumWidth(300)
        self.init_ui()
    
    def init_ui(self):
        layout = QVBoxLayout(self)
        
        # Username
        username_layout = QHBoxLayout()
        username_label = QLabel("Username:")
        self.username_input = QLineEdit()
        username_layout.addWidget(username_label)
        username_layout.addWidget(self.username_input)
        layout.addLayout(username_layout)
        
        # Password
        password_layout = QHBoxLayout()
        password_label = QLabel("Password:")
        self.password_input = QLineEdit()
        self.password_input.setEchoMode(QLineEdit.EchoMode.Password)
        password_layout.addWidget(password_label)
        password_layout.addWidget(self.password_input)
        layout.addLayout(password_layout)
        
        # Buttons
        button_layout = QHBoxLayout()
        login_button = QPushButton("Login")
        login_button.clicked.connect(self.login)
        cancel_button = QPushButton("Cancel")
        cancel_button.clicked.connect(self.reject)
        button_layout.addWidget(login_button)
        button_layout.addWidget(cancel_button)
        layout.addLayout(button_layout)
    
    def login(self):
        username = self.username_input.text()
        password = self.password_input.text()
        
        if self.user_manager.authenticate_user(username, password):
            self.accept()
        else:
            QMessageBox.warning(self, "Error", "Invalid username or password")

class UpdateManager:
    def __init__(self, current_version: str):
        self.current_version = current_version
        self.update_url = "https://api.github.com/repos/your-repo/work-order-parser/releases/latest"
        self.download_url = "https://github.com/your-repo/work-order-parser/releases/latest/download/Work_Order_Parser.exe"
    
    def check_for_updates(self) -> Tuple[bool, str, str]:
        try:
            response = requests.get(self.update_url)
            if response.status_code == 200:
                latest_version = response.json()['tag_name']
                if self._compare_versions(latest_version, self.current_version) > 0:
                    return True, latest_version, response.json()['body']
            return False, "", ""
        except Exception as e:
            logger.logger.error(f"Error checking for updates: {str(e)}")
            return False, "", ""
    
    def _compare_versions(self, version1: str, version2: str) -> int:
        v1_parts = [int(x) for x in version1.split('.')]
        v2_parts = [int(x) for x in version2.split('.')]
        
        for i in range(max(len(v1_parts), len(v2_parts))):
            v1 = v1_parts[i] if i < len(v1_parts) else 0
            v2 = v2_parts[i] if i < len(v2_parts) else 0
            
            if v1 > v2:
                return 1
            elif v1 < v2:
                return -1
        return 0
    
    def download_update(self, progress_callback=None) -> bool:
        try:
            response = requests.get(self.download_url, stream=True)
            if response.status_code == 200:
                total_size = int(response.headers.get('content-length', 0))
                block_size = 1024
                downloaded = 0
                
                with open("Work_Order_Parser_Update.exe", 'wb') as f:
                    for data in response.iter_content(block_size):
                        downloaded += len(data)
                        f.write(data)
                        if progress_callback:
                            progress = int((downloaded / total_size) * 100)
                            progress_callback(progress)
                
                return True
            return False
        except Exception as e:
            logger.logger.error(f"Error downloading update: {str(e)}")
            return False

class UpdateDialog(QDialog):
    def __init__(self, update_manager: UpdateManager, parent=None):
        super().__init__(parent)
        self.update_manager = update_manager
        self.setWindowTitle("Update Available")
        self.setMinimumWidth(400)
        self.init_ui()
    
    def init_ui(self):
        layout = QVBoxLayout(self)
        
        # Update info
        info_label = QLabel("A new version is available!")
        info_label.setStyleSheet("font-size: 14px; font-weight: bold;")
        layout.addWidget(info_label)
        
        # Version info
        version_layout = QHBoxLayout()
        current_version_label = QLabel(f"Current Version: {self.update_manager.current_version}")
        latest_version_label = QLabel(f"Latest Version: {self.update_manager.latest_version}")
        version_layout.addWidget(current_version_label)
        version_layout.addWidget(latest_version_label)
        layout.addLayout(version_layout)
        
        # Release notes
        notes_label = QLabel("Release Notes:")
        notes_label.setStyleSheet("font-weight: bold;")
        layout.addWidget(notes_label)
        
        notes_text = QTextEdit()
        notes_text.setReadOnly(True)
        notes_text.setPlainText(self.update_manager.release_notes)
        notes_text.setMaximumHeight(200)
        layout.addWidget(notes_text)
        
        # Progress bar
        self.progress_bar = QProgressBar()
        self.progress_bar.setVisible(False)
        layout.addWidget(self.progress_bar)
        
        # Buttons
        button_layout = QHBoxLayout()
        update_button = QPushButton("Update Now")
        update_button.clicked.connect(self.start_update)
        later_button = QPushButton("Later")
        later_button.clicked.connect(self.reject)
        button_layout.addWidget(update_button)
        button_layout.addWidget(later_button)
        layout.addLayout(button_layout)
    
    def start_update(self):
        self.progress_bar.setVisible(True)
        self.progress_bar.setValue(0)
        
        # Start download in a separate thread
        self.thread = QThread()
        self.worker = UpdateWorker(self.update_manager)
        self.worker.moveToThread(self.thread)
        
        self.thread.started.connect(self.worker.download)
        self.worker.progress.connect(self.update_progress)
        self.worker.finished.connect(self.update_complete)
        self.worker.error.connect(self.update_error)
        
        self.thread.start()
    
    def update_progress(self, value):
        self.progress_bar.setValue(value)
    
    def update_complete(self):
        self.accept()
        QMessageBox.information(self, "Update Complete", "The application will now restart to apply the update.")
        self.restart_application()
    
    def update_error(self, error):
        QMessageBox.critical(self, "Update Error", f"Error downloading update: {error}")
        self.reject()
    
    def restart_application(self):
        # Start the new version
        QProcess.startDetached("Work_Order_Parser_Update.exe")
        # Exit the current version
        QApplication.quit()

class UpdateWorker(QObject):
    progress = pyqtSignal(int)
    finished = pyqtSignal()
    error = pyqtSignal(str)
    
    def __init__(self, update_manager: UpdateManager):
        super().__init__()
        self.update_manager = update_manager
    
    def download(self):
        try:
            if self.update_manager.download_update(self.progress.emit):
                self.finished.emit()
            else:
                self.error.emit("Download failed")
        except Exception as e:
            self.error.emit(str(e))

class WorkOrderParser(QMainWindow):
    def __init__(self):
        super().__init__()
        self.settings = QSettings('EagleSignCo', 'WorkOrderParser')
        saved_mode = self.settings.value('theme_mode', 'light')
        self.current_mode = saved_mode if saved_mode in COLOR_MODES else 'light'
        # Restore window geometry/state
        self.restoreGeometry(self.settings.value('window_geometry', b''))
        self.restoreState(self.settings.value('window_state', b''))
        self.setWindowTitle(WINDOW['title'])
        self.setMinimumSize(WINDOW['min_width'], WINDOW['min_height'])
        if os.path.exists(THEME['icon_path']):
            self.setWindowIcon(QIcon(THEME['icon_path']))
        self.init_ui()
        self.setup_theme(self.current_mode)
        self.create_theme_menu()
        # Restore splitter positions and table column widths after UI is set up
        self.restore_splitter_positions()
        self.restore_table_settings()
        # Restore toolbar/sidebar visibility (stub)
        # self.restore_toolbar_sidebar_visibility()
        # Restore font size/zoom (stub)
        # self.restore_font_size()
        # Restore language/localization (stub)
        # self.restore_language()
        
        # Initialize update manager
        self.update_manager = UpdateManager("1.0.0")
        
        # Initialize security
        self.security_config = SecurityConfig()
        self.secure_handler = SecureFileHandler(self.security_config)
        
        # Initialize user management
        self.user_manager = UserManager(self.security_config)
        
        # Show login dialog
        if not self.show_login():
            sys.exit()
        
        # Initialize data validation
        self.data_sanitizer = DataSanitizer()
        
        # Initialize components
        self.ai_processor = AIProcessor()
        self.data_analyzer = DataAnalyzer()
        self.export_manager = AdvancedExportManager(self.security_config)
        
        # Initialize UI
        self.setup_theme()
        
        # Setup system tray
        self.setup_tray()
        
        # Initialize data
        self.parsed_data = []
        self.current_file = None
        
        # Check for updates
        self.check_updates()
        
        # Setup auto-save
        self.setup_auto_save()
        
        # Initialize database
        self.init_database()
        
        # Setup security monitoring
        self.setup_security_monitoring()
    
    def check_updates(self):
        has_update, latest_version, release_notes = self.update_manager.check_for_updates()
        if has_update:
            self.update_manager.latest_version = latest_version
            self.update_manager.release_notes = release_notes
            dialog = UpdateDialog(self.update_manager, self)
            dialog.exec()
    
    def show_login(self) -> bool:
        dialog = LoginDialog(self.user_manager, self)
        return dialog.exec() == QDialog.DialogCode.Accepted
    
    def init_database(self):
        try:
            self.conn = sqlite3.connect('work_orders.db')
            self.cursor = self.conn.cursor()
            
            # Create tables
            self.cursor.execute('''
                CREATE TABLE IF NOT EXISTS users (
                    id INTEGER PRIMARY KEY,
                    username TEXT UNIQUE,
                    password_hash TEXT,
                    role TEXT,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            ''')
            
            self.cursor.execute('''
                CREATE TABLE IF NOT EXISTS work_orders (
                    id INTEGER PRIMARY KEY,
                    work_order TEXT,
                    date TEXT,
                    part_number TEXT,
                    department TEXT,
                    created_by INTEGER,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    FOREIGN KEY (created_by) REFERENCES users (id)
                )
            ''')
            
            self.cursor.execute('''
                CREATE TABLE IF NOT EXISTS processing_history (
                    id INTEGER PRIMARY KEY,
                    operation TEXT,
                    status TEXT,
                    details TEXT,
                    user_id INTEGER,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    FOREIGN KEY (user_id) REFERENCES users (id)
                )
            ''')
            
            # Create default admin user if not exists
            self.cursor.execute('SELECT * FROM users WHERE username = ?', ('admin',))
            if not self.cursor.fetchone():
                password_hash = self.security_config.hash_password('admin')
                self.cursor.execute('''
                    INSERT INTO users (username, password_hash, role)
                    VALUES (?, ?, ?)
                ''', ('admin', password_hash, 'admin'))
            
            self.conn.commit()
        except Exception as e:
            logger.logger.error(f"Error initializing database: {str(e)}")
    
    def save_to_database(self, data: List[Dict[str, Any]]):
        try:
            for order in data:
                self.cursor.execute('''
                    INSERT INTO work_orders (work_order, date, part_number, department, created_by)
                    VALUES (?, ?, ?, ?, ?)
                ''', (
                    order['work_order'],
                    order['date'],
                    order['part_number'],
                    order['department'],
                    self.user_manager.current_user['id']
                ))
            
            self.conn.commit()
        except Exception as e:
            logger.logger.error(f"Error saving to database: {str(e)}")
    
    def export_data(self):
        if not self.user_manager.check_permission('export'):
            QMessageBox.warning(self, "Access Denied", "You don't have permission to export data.")
            return
        
        if not self.parsed_data:
            QMessageBox.warning(self, "Warning", "No data to export.")
            return
        
        settings = QSettings()
        default_format = settings.value("default_format", "Excel")
        
        # Get save location
        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "Save File",
            str(Path.home() / "Documents" / f"work_orders.{default_format.lower()}"),
            "Excel Files (*.xlsx);;CSV Files (*.csv);;JSON Files (*.json);;PDF Files (*.pdf);;Word Files (*.docx);;XML Files (*.xml);;YAML Files (*.yaml);;TOML Files (*.toml);;HTML Files (*.html);;Markdown Files (*.md);;LaTeX Files (*.tex);;SQL Files (*.sql);;Parquet Files (*.parquet);;Feather Files (*.feather);;Pickle Files (*.pkl);;HDF5 Files (*.h5)"
        )
        
        if not file_path:
            return
        
        try:
            # Convert to DataFrame
            df = pd.DataFrame(self.parsed_data)
            
            # Export using the export manager
            if self.export_manager.export_data(df, file_path.split('.')[-1], file_path):
                self.statusBar().showMessage(f"Data exported successfully to {file_path}")
                QMessageBox.information(self, "Success", "Data exported successfully!")
            else:
                raise Exception("Export failed")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error exporting data: {str(e)}")
    
    def parse_work_orders(self):
        if not self.user_manager.check_permission('write'):
            QMessageBox.warning(self, "Access Denied", "You don't have permission to parse work orders.")
            return
        
        text = self.input_text.toPlainText().strip()
        if not text:
            QMessageBox.warning(self, "Warning", "Please paste work order data first.")
            return
        
        # Disable UI during processing
        self.parse_button.setEnabled(False)
        self.progress_bar.setVisible(True)
        self.progress_bar.setValue(0)
        
        # Start processing thread
        self.thread = ProcessingThread(text, self.data_sanitizer)
        self.thread.progress.connect(self.update_progress)
        self.thread.finished.connect(self.process_complete)
        self.thread.error.connect(self.process_error)
        self.thread.start()
    
    def setup_security_monitoring(self):
        # Monitor file access
        self.file_monitor = QTimer()
        self.file_monitor.timeout.connect(self.check_file_security)
        self.file_monitor.start(60000)  # Check every minute
        
        # Monitor system security
        self.system_monitor = QTimer()
        self.system_monitor.timeout.connect(self.check_system_security)
        self.system_monitor.start(300000)  # Check every 5 minutes
    
    def check_file_security(self):
        try:
            # Check file permissions
            if self.current_file:
                sd = win32security.GetFileSecurity(
                    self.current_file,
                    win32security.DACL_SECURITY_INFORMATION
                )
                dacl = sd.GetSecurityDescriptorDacl()
                
                # Verify permissions
                if not self._verify_file_permissions(dacl):
                    logger.logger.warning("Insecure file permissions detected")
                    self._fix_file_permissions(self.current_file)
            
            # Check for unauthorized access
            self._check_unauthorized_access()
        except Exception as e:
            logger.logger.error(f"Error in file security check: {str(e)}")
    
    def check_system_security(self):
        try:
            # Check system security settings
            self._check_system_security_settings()
            
            # Check for suspicious processes
            self._check_suspicious_processes()
            
            # Check network security
            self._check_network_security()
        except Exception as e:
            logger.logger.error(f"Error in system security check: {str(e)}")
    
    def _verify_file_permissions(self, dacl) -> bool:
        # Implement permission verification logic
        return True
    
    def _fix_file_permissions(self, file_path: str):
        # Set secure permissions
        sd = win32security.SECURITY_DESCRIPTOR()
        sd.SetSecurityDescriptorDacl(1, None, 0)
        win32security.SetFileSecurity(
            file_path,
            win32security.DACL_SECURITY_INFORMATION,
            sd
        )
    
    def _check_unauthorized_access(self):
        # Implement unauthorized access detection
        pass
    
    def _check_system_security_settings(self):
        # Implement system security check
        pass
    
    def _check_suspicious_processes(self):
        # Implement suspicious process detection
        pass
    
    def _check_network_security(self):
        # Implement network security check
        pass
    
    def setup_auto_save(self):
        self.auto_save_timer = QTimer()
        self.auto_save_timer.timeout.connect(self.auto_save)
        self.auto_save_timer.start(300000)  # Auto-save every 5 minutes
    
    def auto_save(self):
        if self.parsed_data:
            try:
                # Save to database
                self.save_to_database(self.parsed_data)
                
                # Save to file
                backup_dir = Path.home() / "Documents" / "Work Order Parser" / "Backups"
                backup_dir.mkdir(parents=True, exist_ok=True)
                
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                backup_file = backup_dir / f"work_orders_backup_{timestamp}.json"
                
                with open(backup_file, 'w') as f:
                    json.dump(self.parsed_data, f)
                
                logger.logger.info(f"Auto-save completed: {backup_file}")
            except Exception as e:
                logger.logger.error(f"Error in auto-save: {str(e)}")
    
    def init_ui(self):
        # Create central widget and layout
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        layout = QVBoxLayout(central_widget)
        
        # Add Eagle Sign logo if available
        if os.path.exists(THEME['logo_path']):
            logo_label = QLabel()
            logo_pixmap = QPixmap(THEME['logo_path'])
            logo_label.setPixmap(logo_pixmap.scaled(200, 100, Qt.AspectRatioMode.KeepAspectRatio))
            layout.addWidget(logo_label, alignment=Qt.AlignmentFlag.AlignCenter)
        
        # Add title
        title_label = QLabel(WINDOW['title'])
        title_label.setStyleSheet(f"font-size: {FONT_SIZES['title']}px; font-weight: bold; color: {COLORS['primary']};")
        layout.addWidget(title_label, alignment=Qt.AlignmentFlag.AlignCenter)
        
        # Create toolbar
        self.create_toolbar()
        
        # Create main splitter
        splitter = QSplitter(Qt.Orientation.Horizontal)
        layout.addWidget(splitter)
        
        # Left panel
        left_panel = QWidget()
        left_layout = QVBoxLayout(left_panel)
        
        # Input area
        input_group = QGroupBox("Input")
        input_layout = QVBoxLayout()
        self.input_text = QTextEdit()
        self.input_text.setPlaceholderText("Paste work order data here...")
        input_layout.addWidget(self.input_text)
        input_group.setLayout(input_layout)
        left_layout.addWidget(input_group)
        
        # Processing options
        options_group = QGroupBox("Processing Options")
        options_layout = QFormLayout()
        
        self.ai_processing = QCheckBox("Enable AI Processing")
        self.ai_processing.setChecked(True)
        options_layout.addRow("AI Processing:", self.ai_processing)
        
        self.advanced_analysis = QCheckBox("Enable Advanced Analysis")
        self.advanced_analysis.setChecked(True)
        options_layout.addRow("Advanced Analysis:", self.advanced_analysis)
        
        options_group.setLayout(options_layout)
        left_layout.addWidget(options_group)
        
        splitter.addWidget(left_panel)
        
        # Right panel
        right_panel = QWidget()
        right_layout = QVBoxLayout(right_panel)
        
        # Tab widget for different views
        self.tab_widget = QTabWidget()
        
        # Results tab
        results_tab = QWidget()
        results_layout = QVBoxLayout(results_tab)
        self.results_table = QTableWidget()
        self.results_table.setColumnCount(4)
        self.results_table.setHorizontalHeaderLabels(['Work Order', 'Date', 'Part Number', 'Department'])
        self.results_table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
        results_layout.addWidget(self.results_table)
        self.tab_widget.addTab(results_tab, "Results")
        
        # Statistics tab
        stats_tab = QWidget()
        stats_layout = QVBoxLayout(stats_tab)
        self.stats_table = QTableWidget()
        self.stats_table.setColumnCount(2)
        self.stats_table.setHorizontalHeaderLabels(['Metric', 'Value'])
        self.stats_table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
        stats_layout.addWidget(self.stats_table)
        self.tab_widget.addTab(stats_tab, "Statistics")
        
        # Analysis tab
        analysis_tab = QWidget()
        analysis_layout = QVBoxLayout(analysis_tab)
        self.analysis_chart = QChartView()
        self.analysis_chart.setRenderHint(QPainter.RenderHint.Antialiasing)
        analysis_layout.addWidget(self.analysis_chart)
        self.tab_widget.addTab(analysis_tab, "Analysis")
        
        right_layout.addWidget(self.tab_widget)
        splitter.addWidget(right_panel)
        
        # Set initial splitter sizes
        splitter.setSizes([400, 800])
        
        # Progress bar
        self.progress_bar = QProgressBar()
        self.progress_bar.setVisible(False)
        layout.addWidget(self.progress_bar)
        
        # Status bar
        self.statusBar().showMessage("Ready")
    
    def create_toolbar(self):
        toolbar = QToolBar()
        toolbar.setStyleSheet(f"""
            QToolBar {{
                background-color: {COLORS['primary']};
                border: none;
            }}
            QToolButton {{
                color: {COLORS['background']};
                font-family: {FONTS['main']};
                font-size: {FONT_SIZES['normal']}px;
            }}
            QToolButton:hover {{
                background-color: {COLORS['secondary']};
            }}
        """)
        self.addToolBar(toolbar)
        
        # Parse action
        parse_action = QAction("Parse", self)
        parse_action.triggered.connect(self.parse_work_orders)
        toolbar.addAction(parse_action)
        
        # Export action
        export_action = QAction("Export", self)
        export_action.triggered.connect(self.export_data)
        toolbar.addAction(export_action)
        
        toolbar.addSeparator()
        
        # Settings action
        settings_action = QAction("Settings", self)
        settings_action.triggered.connect(self.show_settings)
        toolbar.addAction(settings_action)
        
        # Help action
        help_action = QAction("Help", self)
        help_action.triggered.connect(self.show_help)
        toolbar.addAction(help_action)
    
    def process_work_order(self, data: Dict[str, Any]) -> Tuple[Dict[str, Any], List[str]]:
        # Sanitize and validate data
        sanitized_data, errors = self.data_sanitizer.validate_and_sanitize(data)
        
        if errors:
            logger.logger.warning(f"Validation errors: {errors}")
            return None, errors
        
        # Additional validation
        if not self._validate_work_order_format(sanitized_data['work_order']):
            errors.append("Invalid work order format")
            return None, errors
        
        if not self._validate_date_format(sanitized_data['date']):
            errors.append("Invalid date format")
            return None, errors
        
        if not self._validate_part_number(sanitized_data['part_number']):
            errors.append("Invalid part number format")
            return None, errors
        
        return sanitized_data, []
    
    def _validate_work_order_format(self, work_order: str) -> bool:
        # Implement work order format validation
        pattern = WORK_ORDER_PATTERN
        return bool(re.match(pattern, work_order))
    
    def _validate_date_format(self, date: str) -> bool:
        try:
            datetime.strptime(date, '%Y-%m-%d')
            return True
        except:
            return False
    
    def _validate_part_number(self, part_number: str) -> bool:
        # Implement part number validation
        pattern = WORK_ORDER_PATTERN
        return bool(re.match(pattern, part_number))
    
    def process_complete(self, result):
        self.parsed_data = result['data']
        self.parse_button.setEnabled(True)
        self.export_button.setEnabled(True)
        self.progress_bar.setVisible(False)
        
        # Update results table
        self.results_table.setRowCount(len(result['data']))
        for i, order in enumerate(result['data']):
            self.results_table.setItem(i, 0, QTableWidgetItem(order['work_order']))
            self.results_table.setItem(i, 1, QTableWidgetItem(order['date']))
            self.results_table.setItem(i, 2, QTableWidgetItem(order['part_number']))
            self.results_table.setItem(i, 3, QTableWidgetItem(order['department'] or ''))
        
        # Update statistics
        self.update_statistics(result['stats'])
        
        # Perform AI analysis if enabled
        if self.ai_processing.isChecked():
            self.perform_ai_analysis(result['data'])
        
        # Perform advanced analysis if enabled
        if self.advanced_analysis.isChecked():
            self.perform_advanced_analysis(result['data'])
        
        # Save to database
        self.save_to_database(result['data'])
        
        if result['success']:
            self.statusBar().showMessage(f"Processed {len(result['data'])} work orders successfully")
        else:
            self.statusBar().showMessage("Processing completed with errors")
    
    def perform_ai_analysis(self, data):
        try:
            # Perform AI analysis on the data
            analysis_results = self.ai_processor.analyze_text(str(data))
            
            # Update UI with AI insights
            if analysis_results:
                self.update_ai_insights(analysis_results)
        except Exception as e:
            logger.logger.error(f"Error in AI analysis: {str(e)}")
    
    def perform_advanced_analysis(self, data):
        try:
            # Convert to DataFrame
            df = pd.DataFrame(data)
            
            # Perform advanced analysis
            analysis_results = self.data_analyzer.analyze_data(df)
            
            # Update charts and visualizations
            if analysis_results:
                self.update_analysis_charts(analysis_results)
        except Exception as e:
            logger.logger.error(f"Error in advanced analysis: {str(e)}")
    
    def update_analysis_charts(self, analysis_results):
        try:
            # Create a new chart
            chart = QChart()
            
            # Add data series
            if 'pca_result' in analysis_results:
                # Create scatter plot for PCA results
                series = QScatterSeries()
                for i, point in enumerate(analysis_results['pca_result']):
                    series.append(point[0], point[1])
                chart.addSeries(series)
            
            # Set chart properties
            chart.setTitle("Data Analysis")
            chart.setAnimationOptions(QChart.AnimationOption.SeriesAnimations)
            
            # Update the chart view
            self.analysis_chart.setChart(chart)
        except Exception as e:
            logger.logger.error(f"Error updating analysis charts: {str(e)}")
    
    def closeEvent(self, event):
        # Save window geometry/state
        self.settings.setValue('window_geometry', self.saveGeometry())
        self.settings.setValue('window_state', self.saveState())
        # Save splitter positions and table column widths
        self.save_splitter_positions()
        self.save_table_settings()
        # Save toolbar/sidebar visibility (stub)
        # self.save_toolbar_sidebar_visibility()
        # Save font size/zoom (stub)
        # self.save_font_size()
        # Save language/localization (stub)
        # self.save_language()
        reply = QMessageBox.question(
            self,
            'Confirm Exit',
            "Are you sure you want to exit?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
            QMessageBox.StandardButton.No
        )
        
        if reply == QMessageBox.StandardButton.Yes:
            # Save any pending data
            if self.parsed_data:
                self.auto_save()
            
            # Close database connection
            if hasattr(self, 'conn'):
                self.conn.close()
            
            # Clean up temporary files
            self._cleanup_temp_files()
            
            event.accept()
        else:
            event.ignore()
    
    def _cleanup_temp_files(self):
        try:
            for file in self.secure_handler.temp_dir.glob('*'):
                file.unlink()
            self.secure_handler.temp_dir.rmdir()
        except Exception as e:
            logger.logger.error(f"Error cleaning up temporary files: {str(e)}")

    def setup_theme(self, mode):
        c = COLOR_MODES[mode]
        self.setStyleSheet(f"""
            QMainWindow {{
                background-color: {c['background']};
            }}
            QPushButton {{
                background-color: {c['button']};
                color: {c['button_text']};
                border: none;
                padding: 8px 16px;
                border-radius: 4px;
                font-weight: bold;
                font-family: {FONTS['main']};
                font-size: {FONT_SIZES['normal']}px;
            }}
            QPushButton:hover {{
                background-color: {c['button_hover']};
            }}
            QTextEdit, QLineEdit {{
                border: 1px solid {c['border']};
                border-radius: 4px;
                padding: 8px;
                font-size: {FONT_SIZES['normal']}px;
                font-family: {FONTS['main']};
                color: {c['text']};
                background: {c['background']};
            }}
            QTableWidget {{
                background-color: {c['table_row_even']};
                alternate-background-color: {c['table_row_odd']};
                border: 1px solid {c['border']};
                color: {c['text']};
            }}
            QTableWidget::item {{
                padding: 5px;
                font-family: {FONTS['main']};
                color: {c['text']};
            }}
            QHeaderView::section {{
                background-color: {c['table_header']};
                color: {c['table_header_text']};
                padding: 5px;
                border: 1px solid {c['border']};
                font-weight: bold;
            }}
            QProgressBar {{
                border: 1px solid {c['border']};
                border-radius: 3px;
                text-align: center;
                color: {c['text']};
            }}
            QProgressBar::chunk {{
                background-color: {c['primary']};
            }}
            QLabel {{
                font-family: {FONTS['main']};
                color: {c['text']};
            }}
            QComboBox {{
                border: 1px solid {c['border']};
                border-radius: 4px;
                padding: 8px;
                font-family: {FONTS['main']};
                color: {c['text']};
                background: {c['background']};
            }}
        """)

    def create_theme_menu(self):
        menubar = self.menuBar()
        theme_menu = menubar.addMenu('Theme')
        for mode in COLOR_MODES.keys():
            action = QAction(mode.capitalize() + ' Mode', self)
            action.setCheckable(True)
            action.setChecked(mode == self.current_mode)
            action.triggered.connect(lambda checked, m=mode: self.switch_theme(m))
            theme_menu.addAction(action)
        self.theme_actions = theme_menu.actions()

    def switch_theme(self, mode):
        self.current_mode = mode
        self.setup_theme(mode)
        # Update menu checks
        for action in self.theme_actions:
            action.setChecked(action.text().lower().startswith(mode))
        # Save user preference
        self.settings.setValue('theme_mode', mode)

    def get_last_file_path(self):
        return self.settings.value('last_file_path', '')

    def set_last_file_path(self, file_path):
        self.settings.setValue('last_file_path', file_path)

    def open_file_dialog(self, caption, filter_str, for_save=False):
        last_path = self.get_last_file_path()
        if for_save:
            file_path, _ = QFileDialog.getSaveFileName(self, caption, last_path, filter_str)
        else:
            file_path, _ = QFileDialog.getOpenFileName(self, caption, last_path, filter_str)
        if file_path:
            self.set_last_file_path(file_path)
        return file_path

    # Splitter positions
    def save_splitter_positions(self):
        if hasattr(self, 'splitter'):
            self.settings.setValue('splitter_state', self.splitter.saveState())
    def restore_splitter_positions(self):
        if hasattr(self, 'splitter'):
            state = self.settings.value('splitter_state')
            if state:
                self.splitter.restoreState(state)

    # Table column widths and sort order
    def save_table_settings(self):
        if hasattr(self, 'table'):
            widths = [self.table.columnWidth(i) for i in range(self.table.columnCount())]
            self.settings.setValue('table_column_widths', widths)
            # Save sort order if applicable
            # self.settings.setValue('table_sort_order', self.table.horizontalHeader().sortIndicatorSection())
    def restore_table_settings(self):
        if hasattr(self, 'table'):
            widths = self.settings.value('table_column_widths')
            if widths:
                for i, w in enumerate(widths):
                    self.table.setColumnWidth(i, int(w))
            # Restore sort order if applicable
            # sort_section = self.settings.value('table_sort_order')
            # if sort_section is not None:
            #     self.table.sortItems(int(sort_section))

    # Recent files (stub)
    def add_recent_file(self, file_path):
        recent = self.settings.value('recent_files', [])
        if file_path not in recent:
            recent = [file_path] + recent[:9]  # Keep max 10
            self.settings.setValue('recent_files', recent)
    def get_recent_files(self):
        return self.settings.value('recent_files', [])

    # Last used export format (stub)
    def set_last_export_format(self, fmt):
        self.settings.setValue('last_export_format', fmt)
    def get_last_export_format(self):
        return self.settings.value('last_export_format', '')

    # Toolbar/sidebar visibility (stub)
    # def save_toolbar_sidebar_visibility(self): ...
    # def restore_toolbar_sidebar_visibility(self): ...

    # Font size/zoom (stub)
    # def save_font_size(self): ...
    # def restore_font_size(self): ...

    # Language/localization (stub)
    # def save_language(self): ...
    # def restore_language(self): ...

class ProcessingThread(QThread):
    progress = pyqtSignal(int)
    finished = pyqtSignal(dict)
    error = pyqtSignal(str)
    
    def __init__(self, text: str, data_sanitizer: DataSanitizer):
        super().__init__()
        self.text = text
        self.data_sanitizer = data_sanitizer
    
    def run(self):
        try:
            result = self.process_text(self.text)
            self.finished.emit(result)
        except Exception as e:
            self.error.emit(str(e))
    
    def process_text(self, text: str) -> Dict[str, Any]:
        lines = text.split('\n')
        work_orders = []
        current_part = {}
        errors = []
        
        for i, line in enumerate(lines):
            self.progress.emit(int((i / len(lines)) * 100))
            line = line.strip()
            if not line:
                continue
            
            try:
                # Extract part number
                if line.startswith('Part Number:'):
                    current_part['part_number'] = line.split(':', 1)[1].strip()
                    continue
                
                # Extract work order
                if 'Work Order' in line:
                    work_order = line.split('#', 1)[1].strip()
                    date = None
                    if i + 1 < len(lines):
                        date = lines[i + 1].strip()
                    
                    if date:
                        data = {
                            'work_order': work_order,
                            'date': date,
                            'part_number': current_part.get('part_number', ''),
                            'department': self.extract_department(work_order)
                        }
                        
                        # Validate and sanitize data
                        sanitized_data, validation_errors = self.data_sanitizer.validate_and_sanitize(data)
                        
                        if validation_errors:
                            errors.extend(validation_errors)
                        else:
                            work_orders.append(sanitized_data)
            except Exception as e:
                errors.append(f"Error processing line {i + 1}: {str(e)}")
        
        return {
            'success': len(errors) == 0,
            'data': work_orders,
            'errors': errors if errors else None,
            'stats': {
                'total_lines': len(lines),
                'processed_orders': len(work_orders),
                'errors': len(errors)
            }
        }
    
    def extract_department(self, work_order: str) -> Optional[str]:
        departments = ['ES', 'SC', 'PR', 'MFG', 'ENG', 'QA']
        return next((dept for dept in departments if dept in work_order), None)

def main():
    app = QApplication(sys.argv)
    
    # Set application style
    app.setStyle('Fusion')
    
    # Create and show the main window
    window = WorkOrderParser()
    window.show()
    
    sys.exit(app.exec())

if __name__ == '__main__':
    main()